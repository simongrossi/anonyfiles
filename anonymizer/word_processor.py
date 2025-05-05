# anonyfiles-main/anonymizer/word_processor.py
from docx import Document
import os # Import the os module for directory creation

def extract_text_from_docx(path):
    """
    Extrait le texte d'un fichier .docx, paragraphe par paragraphe.
    """
    doc = Document(path)
    return "\n".join([p.text for p in doc.paragraphs])

def replace_entities_in_docx(path, replacements, output_path):
    """
    Remplace les entités détectées dans un fichier .docx en reconstruisant
    les paragraphes avec le texte anonymisé.
    ATTENTION : Cette méthode va supprimer toute la mise en forme originale
    à l'intérieur des paragraphes modifiés.
    """
    doc = Document(path)

    # On applique les remplacements un par un sur chaque paragraphe
    for p in doc.paragraphs:
        original_text = p.text
        modified_text = original_text

        # Trie les remplacements par longueur décroissante pour gérer les chevauchements potentiels
        # (par exemple, remplacer "Dr. John Smith" avant "John Smith")
        # Note : Ce tri est moins critique avec la méthode simple .replace(), mais reste une bonne pratique.
        sorted_replacements = sorted(replacements.items(), key=lambda item: len(item[0]), reverse=True)


        for original, replacement in sorted_replacements: # Utiliser les remplacements triés
            # Utilise un simple remplacement de chaîne sur le texte du paragraphe.
            # original.strip() permet de mieux gérer les entités détectées avec des espaces/newlines
            # potentiels autour d'elles dans le texte original du paragraphe.
            if original.strip() in modified_text: # Vérifie si l'entité (sans espaces/newlines) est présente
                 modified_text = modified_text.replace(original.strip(), replacement) # Effectue le remplacement

        # Si modification détectée
        if modified_text != original_text:
            # Supprime tous les runs existants du paragraphe
            # On itère sur une copie des éléments de run car on modifie la liste pendant l'itération
            for run_element in p._element.xpath('./w:r'):
                run_element.getparent().remove(run_element)

            # Ajoute le texte modifié dans un nouveau run
            # Ce nouveau run aura le formatage par défaut du paragraphe ou du document
            if modified_text.strip(): # Ajoute un run seulement si le texte modifié n'est pas vide après strip
                p.add_run(modified_text)
            # Si modified_text.strip() est vide, le paragraphe sera vide (sans runs)


    # S'assurer que le dossier de sortie existe (ajout pour robustesse)
    output_dir = os.path.dirname(output_path)
    # Vérifie si le chemin contient une partie répertoire et si elle n'existe pas
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)

    # Sauvegarder le document modifié
    doc.save(output_path)

