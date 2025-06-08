from anonyfiles_cli.anonymizer.word_processor import DocxProcessor
from docx import Document
from pathlib import Path
import tempfile
import pytest

def test_extract_blocks_docx():
    tmp = tempfile.NamedTemporaryFile("w+b", delete=False, suffix=".docx")
    doc = Document()
    doc.add_paragraph("Pierre est à Paris.")
    doc.add_paragraph("Ceci est un test.")
    doc.save(tmp.name)
    processor = DocxProcessor()
    blocks = processor.extract_blocks(tmp.name)
    assert len(blocks) == 2
    assert "Pierre" in blocks[0]

def test_reconstruct_and_write_anonymized_file_docx():
    tmp_in = tempfile.NamedTemporaryFile("w+b", delete=False, suffix=".docx")
    tmp_out = tempfile.NamedTemporaryFile("w+b", delete=False, suffix=".docx")
    doc = Document()
    doc.add_paragraph("Pierre à Paris.")
    doc.save(tmp_in.name)
    processor = DocxProcessor()

    blocks = processor.extract_blocks(tmp_in.name)
    final_blocks = [b.replace("Pierre", "NOM003").replace("Paris", "VILLE_Z") for b in blocks]

    processor.reconstruct_and_write_anonymized_file(
        Path(tmp_out.name),
        final_blocks,
        Path(tmp_in.name),
    )

    doc_result = Document(tmp_out.name)
    assert "NOM003" in doc_result.paragraphs[0].text
    assert "VILLE_Z" in doc_result.paragraphs[0].text


def test_reconstruct_and_write_anonymized_file_docx_preserves_formatting():
    tmp_in = tempfile.NamedTemporaryFile("w+b", delete=False, suffix=".docx")
    tmp_out = tempfile.NamedTemporaryFile("w+b", delete=False, suffix=".docx")

    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("Pierre")
    r1.bold = True
    p.add_run(" à ")
    r3 = p.add_run("Paris")
    r3.italic = True
    doc.save(tmp_in.name)

    processor = DocxProcessor()
    blocks = processor.extract_blocks(tmp_in.name)
    final_blocks = [b.replace("Pierre", "NOM003").replace("Paris", "VILLE_Z") for b in blocks]

    processor.reconstruct_and_write_anonymized_file(
        Path(tmp_out.name),
        final_blocks,
        Path(tmp_in.name),
    )

    doc_result = Document(tmp_out.name)
    runs = doc_result.paragraphs[0].runs
    assert runs[0].text == "NOM003"
    assert runs[0].bold
    assert runs[-1].text.endswith("VILLE_Z")
    assert runs[-1].italic


def test_reconstruct_and_write_anonymized_file_docx_mismatch():
    tmp_in = tempfile.NamedTemporaryFile("w+b", delete=False, suffix=".docx")
    tmp_out = tempfile.NamedTemporaryFile("w+b", delete=False, suffix=".docx")
    doc = Document()
    doc.add_paragraph("Texte unique")
    doc.save(tmp_in.name)

    processor = DocxProcessor()

    # Fournir une liste vide pour provoquer un décalage de compte
    with pytest.raises(ValueError):
        processor.reconstruct_and_write_anonymized_file(
            Path(tmp_out.name),
            [],
            Path(tmp_in.name),
        )
