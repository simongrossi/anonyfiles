# anonymizer/word_processor.py

from pathlib import Path
import logging
from typing import Any, Iterator
from docx import Document
from .base_processor import BaseProcessor
from .type_defs import TextBlocks

logger = logging.getLogger(__name__)


class DocxProcessor(BaseProcessor):
    """
    Processor pour les fichiers .docx.
    - Traverse hiérarchiquement : Paragraphes du corps -> Tableaux (récursifs).
    - Préserve l'intégrité de la structure tout en anonymisant l'ensemble du contenu.
    """

    @staticmethod
    def _open_document(path: Path) -> Any:
        """Ouvre un document .docx en remontant une erreur claire si illisible."""
        try:
            return Document(str(path))
        except Exception as exc:  # noqa: BLE001 - on veut un message unifié
            raise ValueError(
                f"Fichier .docx illisible ou corrompu: {Path(path).name} ({exc})"
            ) from exc

    def _iter_block_items(self, parent_elt: Any) -> Iterator[Any]:
        """
        Générateur récursif qui parcourt tous les éléments contenant du texte
        (Paragraphes directs, et Cellules de Tableaux) dans un ordre déterministe.
        """
        # 1. Paragraphes directs
        if hasattr(parent_elt, "paragraphs"):
            for paragraph in parent_elt.paragraphs:
                yield paragraph

        # 2. Tableaux (qui contiennent des lignes -> cellules -> paragraphes/tables)
        if hasattr(parent_elt, "tables"):
            for table in parent_elt.tables:
                for row in table.rows:
                    for cell in row.cells:
                        # Récursion: une cellule contient des paragraphes et potentiellement d'autres tables
                        yield from self._iter_block_items(cell)

    def _iter_header_footer_parts(self, doc: Any) -> Iterator[Any]:
        """
        Itère sur les en-têtes et pieds de page *propres* à chaque section.

        Les en-têtes/pieds liés à la section précédente (``is_linked_to_previous``)
        ne contiennent pas de définition propre : on les ignore pour éviter de
        traiter deux fois le même contenu (ou des parties vides héritées).
        """
        for section in doc.sections:
            for part in (
                section.header,
                section.footer,
                section.first_page_header,
                section.first_page_footer,
                section.even_page_header,
                section.even_page_footer,
            ):
                if part is None:
                    continue
                if getattr(part, "is_linked_to_previous", False):
                    continue
                yield part

    def _iter_document_paragraphs(self, doc: Any) -> Iterator[Any]:
        """
        Ordre déterministe couvrant TOUT le texte anonymisable d'un document :
        corps (paragraphes + tableaux), puis en-têtes et pieds de page.

        ``extract_blocks`` et ``reconstruct_and_write_anonymized_file`` partagent
        ce parcours pour garantir l'alignement index-par-index des blocs.
        """
        yield from self._iter_block_items(doc)
        for part in self._iter_header_footer_parts(doc):
            yield from self._iter_block_items(part)

    def extract_blocks(self, input_path: Path, **kwargs: Any) -> TextBlocks:
        """
        Extrait TOUS les blocs de texte (Body + Tables + En-têtes/Pieds de page).
        Retourne une liste plate de chaînes de caractères.
        """
        doc = self._open_document(input_path)
        blocks: TextBlocks = []

        for paragraph in self._iter_document_paragraphs(doc):
            blocks.append(paragraph.text)

        return blocks

    def reconstruct_and_write_anonymized_file(
        self,
        output_path: Path,
        final_processed_blocks: TextBlocks,
        original_input_path: Path,
        **kwargs: Any,
    ) -> None:
        """
        Reconstruit le document DOCX en injectant les blocs anonymisés.
        Gère le corps du texte ET les tableaux.
        """
        doc = self._open_document(original_input_path)
        target_paragraphs = list(self._iter_document_paragraphs(doc))

        count_expected = len(target_paragraphs)
        if count_expected != len(final_processed_blocks):
            # Désalignement = bug d'extraction/reconstruction. On échoue de façon
            # explicite plutôt que de produire un document partiellement anonymisé
            # (risque de fuite de données non anonymisées).
            raise ValueError(
                "Mismatch Reconstruct Word: "
                f"{count_expected} paragraphes (corps + tableaux + en-têtes/pieds) "
                f"vs {len(final_processed_blocks)} blocs fournis."
            )

        for i, p in enumerate(target_paragraphs):
            if i >= len(final_processed_blocks):
                break

            new_text = final_processed_blocks[i]

            # Si le texte n'a pas changé, on ne touche à rien (préserve 100% du formatage complexe)
            if p.text == new_text:
                continue

            # Logique de conservation des styles (Run distribution)
            # On essaie de répartir le nouveau texte dans les runs existants
            runs = p.runs
            if not runs:
                if new_text:
                    p.add_run(new_text)
                continue

            # Si on a du texte à écrire
            # Approche simple et robuste pour éviter les corruptions XML :
            # 1. On calcule la longueur totale disponible dans les runs actuels ? Non, le texte change de taille.
            # 2. Approche "Buckets" : on remplit les runs un par un avec le nouveau texte.

            remaining_text = new_text
            for run in runs:
                if not remaining_text:
                    # Plus de texte à mettre, on vide ce run (il devient invisible)
                    run.text = ""
                    continue

                # Décision arbitraire : on essaie de mettre autant de caractères qu'il y en avait ?
                # Ou on met tout dans le premier ?
                # Mettre tout dans le premier est le plus sûr pour la lisibilité,
                # mais "répartir" aide parfois si le mot en gras était au milieu.
                # Vu qu'on anonymise, la structure sémantique change. "Jean Dupont" (Gras) devient "PER_1" (Gras).
                # Le plus simple est de tout mettre dans le run[0] et vider les autres,
                # OU d'utiliser une logique proportionnelle simple.

                # Gardons la logique du code précédent qui semblait vouloir préserver "autant que possible".
                # MAIS pour plus de fiabilité sur la donnée (éviter de couper au milieu d'un run),
                # la méthode la plus sûre pour l'ETL est :
                # - Run 1 reçoit tout le texte (hérite du style du début de phrase)
                # - Runs suivants vidés.

                # SAUF si l'ancien code avait une meilleure logique.
                # L'ancien code faisait :
                # if pointer < len(new_text): runs[-1].text += new_text[pointer:]
                # Reprenons cette logique "bucket" simple qui préserve la longueur visuelle relative :

                run_len = len(run.text)
                # Si le run était vide, on n'y met rien sauf si c'est le dernier et qu'il reste du texte
                if run_len == 0:
                    run.text = ""
                else:
                    # On prend une tranche
                    chunk = remaining_text[:run_len]
                    run.text = chunk
                    remaining_text = remaining_text[run_len:]

            # S'il reste du texte après avoir rempli tous les slots existants, on l'ajoute au dernier run
            if remaining_text and runs:
                runs[-1].text += remaining_text

        output_dir = Path(output_path).parent
        if not output_dir.exists():
            output_dir.mkdir(parents=True, exist_ok=True)

        doc.save(output_path)
