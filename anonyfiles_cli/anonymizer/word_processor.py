# anonymizer/word_processor.py

import os
import logging
from docx import Document
from .base_processor import BaseProcessor
from .utils import apply_positional_replacements

logger = logging.getLogger(__name__)

class DocxProcessor(BaseProcessor):
    """
    Processor pour les fichiers .docx.
    - Chaque paragraphe est traité comme un bloc.
    """

    def extract_blocks(self, input_path):
        """
        Extrait chaque paragraphe du document DOCX comme un bloc de texte.
        Retourne une liste de paragraphes (chaînes).
        """
        doc = Document(input_path)
        return [p.text for p in doc.paragraphs]

    def replace_entities(
        self,
        input_path,
        output_path,
        replacements,
        entities_per_block_with_offsets
    ):
        """
        Remplace les entités dans chaque paragraphe du document DOCX, puis sauvegarde le résultat.
        """
        doc = Document(input_path)
        paragraphs = doc.paragraphs

        if len(paragraphs) != len(entities_per_block_with_offsets):
            # Mismatch logique, sécurité !
            raise ValueError(f"Le nombre de paragraphes ({len(paragraphs)}) ne correspond pas "
                             f"au nombre de listes d'entités fournies ({len(entities_per_block_with_offsets)}).")

        if not paragraphs or all(not p.text.strip() for p in paragraphs):
            doc.save(output_path)
            return

        for i, p in enumerate(paragraphs):
            original_text = p.text
            entities_for_this_paragraph = entities_per_block_with_offsets[i]

            if original_text.strip() and entities_for_this_paragraph:
                anonymized_text = apply_positional_replacements(
                    original_text,
                    replacements,
                    entities_for_this_paragraph
                )
            else:
                anonymized_text = original_text

            # Supprimer tout le contenu du paragraphe (et donc son formatage !)
            for run_element in p._element.xpath('./w:r'):
                run_element.getparent().remove(run_element)

            if anonymized_text.strip():
                p.add_run(anonymized_text)
            # Sinon on laisse le paragraphe vide

        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)

        doc.save(output_path)

    def reconstruct_and_write_anonymized_file(
        self,
        output_path,
        final_processed_blocks,
        original_input_path,
        **kwargs
    ):
        """Reconstruit un document DOCX à partir des blocs traités et l'enregistre."""
        doc = Document(original_input_path)
        paragraphs = doc.paragraphs

        expected_count = len(paragraphs)
        if expected_count != len(final_processed_blocks):
            logger.warning(
                "Mismatch entre %s paragraphes attendus et %s blocs fournis pour %s",
                expected_count,
                len(final_processed_blocks),
                output_path,
            )
            raise ValueError(
                f"Le nombre de paragraphes ({expected_count}) ne correspond pas au nombre de blocs finaux ({len(final_processed_blocks)})."
            )

        for i, p in enumerate(paragraphs):
            new_text = ""
            if i < len(final_processed_blocks):
                new_text = final_processed_blocks[i]

            for run_element in p._element.xpath('./w:r'):
                run_element.getparent().remove(run_element)

            if new_text.strip():
                p.add_run(new_text)

        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)

        doc.save(output_path)
